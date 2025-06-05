"""
data_extractors.py - Text extraction utilities for various file formats
Supports .txt, .pdf, .docx, and other document formats for business info loading
Part of the Blue AI Coacher Bot system.
"""

import io
import os
import logging
from typing import Optional, Union, BinaryIO, List, Dict
import mimetypes
from datetime import datetime

# Logger will be injected from other modules
logger = None

def setup_data_extractors(logger_instance):
    """Initialize this module with the logger"""
    global logger
    logger = logger_instance

def extract_text_from_pdf(file_data: Union[bytes, BinaryIO]) -> str:
    """
    Extract text from PDF files using PyPDF2
    
    Args:
        file_data: PDF file as bytes or file-like object
        
    Returns:
        str: Extracted text content
    """
    try:
        import PyPDF2
        
        if isinstance(file_data, bytes):
            file_data = io.BytesIO(file_data)
        
        pdf_reader = PyPDF2.PdfReader(file_data)
        text_content = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content += f"\n--- صفحه {page_num + 1} ---\n"
                    text_content += page_text + "\n"
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                continue
        
        if not text_content.strip():
            raise Exception("هیچ متنی از فایل PDF استخراج نشد")
            
        logger.info(f"Successfully extracted {len(text_content)} characters from PDF")
        return text_content.strip()
        
    except ImportError:
        logger.error("PyPDF2 library not installed. Install with: pip install PyPDF2")
        raise Exception("کتابخانه PyPDF2 نصب نیست. برای پردازش فایل PDF، این کتابخانه مورد نیاز است.")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise Exception(f"خطا در استخراج متن از فایل PDF: {str(e)}")

def extract_text_from_docx(file_data: Union[bytes, BinaryIO]) -> str:
    """
    Extract text from DOCX files using python-docx
    
    Args:
        file_data: DOCX file as bytes or file-like object
        
    Returns:
        str: Extracted text content
    """
    try:
        from docx import Document
        
        if isinstance(file_data, bytes):
            file_data = io.BytesIO(file_data)
        
        doc = Document(file_data)
        text_content = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content += " | ".join(row_text) + "\n"
        
        if not text_content.strip():
            raise Exception("هیچ متنی از فایل Word استخراج نشد")
            
        logger.info(f"Successfully extracted {len(text_content)} characters from DOCX")
        return text_content.strip()
        
    except ImportError:
        logger.error("python-docx library not installed. Install with: pip install python-docx")
        raise Exception("کتابخانه python-docx نصب نیست. برای پردازش فایل Word، این کتابخانه مورد نیاز است.")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"خطا در استخراج متن از فایل Word: {str(e)}")

def extract_text_from_txt(file_data: Union[bytes, BinaryIO], encoding: str = 'utf-8') -> str:
    """
    Extract text from TXT files with encoding detection
    
    Args:
        file_data: TXT file as bytes or file-like object
        encoding: Text encoding (default: utf-8)
        
    Returns:
        str: Extracted text content
    """
    try:
        if isinstance(file_data, bytes):
            # Try different encodings
            encodings_to_try = ['utf-8', 'utf-16', 'cp1256', 'iso-8859-1', 'windows-1252']
            
            for enc in encodings_to_try:
                try:
                    text_content = file_data.decode(enc)
                    logger.info(f"Successfully decoded TXT file using {enc} encoding")
                    return text_content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text_content = file_data.decode('utf-8', errors='replace')
            logger.warning("Used UTF-8 with error replacement for TXT file")
            return text_content.strip()
        else:
            # File-like object
            text_content = file_data.read()
            if isinstance(text_content, bytes):
                text_content = text_content.decode(encoding, errors='replace')
            return text_content.strip()
            
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise Exception(f"خطا در استخراج متن از فایل متنی: {str(e)}")

def extract_text_from_rtf(file_data: Union[bytes, BinaryIO]) -> str:
    """
    Extract text from RTF files using striprtf
    
    Args:
        file_data: RTF file as bytes or file-like object
        
    Returns:
        str: Extracted text content
    """
    try:
        from striprtf.striprtf import rtf_to_text
        
        if isinstance(file_data, bytes):
            rtf_content = file_data.decode('utf-8', errors='replace')
        else:
            rtf_content = file_data.read()
            if isinstance(rtf_content, bytes):
                rtf_content = rtf_content.decode('utf-8', errors='replace')
        
        text_content = rtf_to_text(rtf_content)
        
        if not text_content.strip():
            raise Exception("هیچ متنی از فایل RTF استخراج نشد")
            
        logger.info(f"Successfully extracted {len(text_content)} characters from RTF")
        return text_content.strip()
        
    except ImportError:
        logger.error("striprtf library not installed. Install with: pip install striprtf")
        raise Exception("کتابخانه striprtf نصب نیست. برای پردازش فایل RTF، این کتابخانه مورد نیاز است.")
    except Exception as e:
        logger.error(f"Error extracting text from RTF: {e}")
        raise Exception(f"خطا در استخراج متن از فایل RTF: {str(e)}")

def detect_file_type(file_name: str, mime_type: Optional[str] = None) -> str:
    """
    Detect file type based on filename and MIME type
    
    Args:
        file_name: Name of the file
        mime_type: MIME type if available
        
    Returns:
        str: Detected file type ('pdf', 'docx', 'txt', 'rtf', 'unknown')
    """
    file_extension = os.path.splitext(file_name.lower())[1]
    
    # Check by extension first
    extension_map = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',  # Treat DOC as DOCX (requires python-docx)
        '.txt': 'txt',
        '.text': 'txt',
        '.rtf': 'rtf'
    }
    
    if file_extension in extension_map:
        return extension_map[file_extension]
    
    # Check by MIME type
    if mime_type:
        mime_map = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'docx',
            'text/plain': 'txt',
            'text/rtf': 'rtf',
            'application/rtf': 'rtf'
        }
        
        if mime_type in mime_map:
            return mime_map[mime_type]
    
    # Fallback: try to guess from filename
    if any(ext in file_name.lower() for ext in ['.pdf']):
        return 'pdf'
    elif any(ext in file_name.lower() for ext in ['.doc', '.docx']):
        return 'docx'
    elif any(ext in file_name.lower() for ext in ['.txt', '.text']):
        return 'txt'
    elif any(ext in file_name.lower() for ext in ['.rtf']):
        return 'rtf'
    
    return 'unknown'

def extract_text_from_file(file_data: Union[bytes, BinaryIO], file_name: str, mime_type: Optional[str] = None) -> str:
    """
    Main function to extract text from various file formats
    
    Args:
        file_data: File content as bytes or file-like object
        file_name: Name of the file
        mime_type: MIME type if available
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If file type is not supported or extraction fails
    """
    file_type = detect_file_type(file_name, mime_type)
    
    logger.info(f"Extracting text from file: {file_name} (detected type: {file_type})")
    
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'txt': extract_text_from_txt,
        'rtf': extract_text_from_rtf
    }
    
    if file_type not in extractors:
        supported_types = ', '.join(extractors.keys())
        raise Exception(f"نوع فایل '{file_type}' پشتیبانی نمی‌شود. فرمت‌های پشتیبانی شده: {supported_types}")
    
    try:
        extractor = extractors[file_type]
        extracted_text = extractor(file_data)
        
        if not extracted_text or not extracted_text.strip():
            raise Exception("هیچ متنی از فایل استخراج نشد یا فایل خالی است")
        
        # Add metadata header
        header = f"=== متن استخراج شده از فایل: {file_name} ===\n"
        header += f"نوع فایل: {file_type.upper()}\n"
        header += f"طول متن: {len(extracted_text)} کاراکتر\n"
        header += "=" * 50 + "\n\n"
        
        return header + extracted_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from {file_name}: {e}")
        raise

def extract_text_from_multiple_files(files_data: List[Dict], combine_header: bool = True) -> str:
    """
    Extract text from multiple files and combine them
    
    Args:
        files_data: List of dictionaries containing file information
                   Each dict should have: 'content', 'name', 'mime_type'
        combine_header: Whether to add a combined header
        
    Returns:
        str: Combined extracted text from all files
        
    Raises:
        Exception: If no files could be processed successfully
    """
    if not files_data:
        raise Exception("لیست فایل‌ها خالی است")
    
    all_extracted_texts = []
    successfully_processed = []
    failed_files = []
    
    for i, file_info in enumerate(files_data):
        try:
            extracted_text = extract_text_from_file(
                file_info['content'], 
                file_info['name'], 
                file_info.get('mime_type')
            )
            
            # Add file separator
            file_header = f"\n{'='*60}\n"
            file_header += f"📄 فایل #{i+1}: {file_info['name']}\n"
            file_header += f"{'='*60}\n\n"
            
            all_extracted_texts.append(file_header + extracted_text)
            successfully_processed.append(file_info['name'])
            
            logger.info(f"Successfully processed file {i+1}/{len(files_data)}: {file_info['name']}")
            
        except Exception as e:
            logger.error(f"Failed to process file {file_info['name']}: {e}")
            failed_files.append(f"{file_info['name']}: {str(e)}")
    
    if not all_extracted_texts:
        error_details = "\n".join([f"• {error}" for error in failed_files])
        raise Exception(f"هیچ فایلی قابل پردازش نبود:\n{error_details}")
    
    # Combine all texts
    combined_text = "\n\n".join(all_extracted_texts)
    
    if combine_header:
        header = f"{'='*80}\n"
        header += f"📚 مجموعه اطلاعات کسب و کار از {len(successfully_processed)} فایل\n"
        header += f"📅 تاریخ پردازش: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        header += f"✅ فایل‌های موفق: {', '.join(successfully_processed)}\n"
        
        if failed_files:
            header += f"❌ فایل‌های ناموفق ({len(failed_files)}): {', '.join([f.split(':')[0] for f in failed_files])}\n"
        
        header += f"📊 مجموع کاراکترها: {len(combined_text):,}\n"
        header += f"{'='*80}\n\n"
        
        combined_text = header + combined_text
    
    logger.info(f"Successfully combined text from {len(successfully_processed)} files. Total length: {len(combined_text)} characters")
    
    return combined_text

def get_supported_file_types() -> dict:
    """
    Get list of supported file types and their descriptions
    
    Returns:
        dict: Mapping of file types to descriptions
    """
    return {
        'pdf': 'فایل‌های PDF',
        'docx': 'فایل‌های Word (DOC/DOCX)',
        'txt': 'فایل‌های متنی (TXT)',
        'rtf': 'فایل‌های RTF'
    }

def validate_file_size(file_size: int, max_size_mb: int = 20) -> bool:
    """
    Validate file size
    
    Args:
        file_size: File size in bytes
        max_size_mb: Maximum allowed size in MB
        
    Returns:
        bool: True if file size is acceptable
    """
    max_size_bytes = max_size_mb * 1024 * 1024
    return file_size <= max_size_bytes

def get_file_processing_help_text() -> str:
    """
    Get help text about file processing capabilities for regular chat
    
    Returns:
        str: Help text explaining file processing in chat
    """
    supported_types = get_supported_file_types()
    types_list = "\n".join([f"📄 {desc}" for desc in supported_types.values()])
    
    return f"""📁 **پردازش فایل در چت**

می‌توانید فایل‌های مختلف را مستقیماً در چت ارسال کنید و ربات محتوای آن‌ها را استخراج کرده و تحلیل می‌کند.

**فرمت‌های پشتیبانی شده:**
{types_list}

**نحوه استفاده:**
1️⃣ فایل خود را در چت ارسال کنید
2️⃣ اختیاری: توضیحی درباره آنچه می‌خواهید اضافه کنید
3️⃣ ربات محتوای فایل را استخراج می‌کند
4️⃣ تحلیل و پاسخ مناسب ارائه می‌دهد

**مثال:**
"این قرارداد را بررسی کن و نکات مهم را برایم خلاصه کن"
[فایل PDF قرارداد ضمیمه]

**محدودیت‌ها:**
• حداکثر حجم فایل: ۲۰ مگابایت
• فایل‌های رمزدار پشتیبانی نمی‌شوند"""

# Installation instructions for required libraries
INSTALLATION_INSTRUCTIONS = """
برای استفاده از امکان استخراج متن از فایل‌ها، کتابخانه‌های زیر باید نصب باشند:

pip install PyPDF2 python-docx striprtf

PyPDF2: برای فایل‌های PDF
python-docx: برای فایل‌های Word
striprtf: برای فایل‌های RTF
"""
